#THƯ VIỆN
import openai
from TikTokLive import TikTokLiveClient
from TikTokLive.types.events import CommentEvent
import asyncio
import csv
import pandas as pd
import gspread
from docx import Document
import os 
import threading
import datetime
from pathlib import Path
import time
import pygame
def get_audio_time(file_path):
    pygame.mixer.init()
    sound = pygame.mixer.Sound(file_path)
    duration_in_seconds = sound.get_length()
    pygame.mixer.quit()
    return duration_in_seconds
def url(name_folder_or_file):
    url_folder = Path(__file__).parent
    file_path=os.path.join(url_folder, name_folder_or_file)
    return file_path
list_data=[]
#XÂY DỰNG BỘ CHUYỂN TEENCODE SANG TIẾNG VIỆT
df_teencode=pd.read_csv(url(r'teencode.csv'),names=['teencode','vietnamese'])
df_teencode_size=len(df_teencode)
##đổi 1 từ teencode thành tiếng việt
def one_tencode_to_VN(text):
    for i in range(df_teencode_size):
        if text == df_teencode['teencode'][i]:
            return df_teencode['vietnamese'][i]
##đổi một câu có chứa teencode thành tiếng việt       
def tencode_to_VN(text):
    str=' '
    text_process=[]
    text=text.lower()
    cups=text.split()
    for cup in cups:
        VN= one_tencode_to_VN(cup)
        if VN != None:
            text_process.append(VN)
        else :
            text_process.append(cup)
    return str.join(text_process)

#XÂY DỰNG BỘ LỌC NHỮNG KÍ TỰ KHÔNG THUỘC MÃ HÓA UTF-8
##xóa những kí tự không thuộc mã hóa UTF-8 ra khỏi câu
def remove_non_UTF8(text):
    words=text.split()
    valid_words=[word for word in words if word.encode('utf-8')]
    return ' '.join(valid_words)
 
#XÂY DỰNG BỘ TRÍCH SUẤT COMMENT, XỬ LÝ COMMENT VÀ LƯU COMMENT VÀO FILE 'data.csv'
##crawling comment
async def crawling(id):
    id='@'+str(id)
    global list_data
    # Instantiate the client with the user's username
    client: TikTokLiveClient = TikTokLiveClient(unique_id=id)
    # Define how you want to handle specific events via decorator
    @client.on("comment")
    # Notice no decorator?
    async def on_comment(event: CommentEvent):
    # Mở tệp CSV hiện có với chế độ ghi
            #trích suất tên người dùng
            user=event.user.nickname
            #trích suất nội dung comment
            comment=event.comment
            #chuẩn hóa comment
            comment=tencode_to_VN(comment)
            #Loại bỏ những kí tự không nằm trong UTF-8
            comment=remove_non_UTF8(comment)
            user=remove_non_UTF8(user)
            print(comment)
            data_to_append = [user,comment]
            list_data.append(data_to_append)
            # Thêm dữ liệu vào tệp CSV
    # Define handling an event via a "callback"
    client.add_listener("comment", on_comment)
    await client.start()

##Model_Check
#XÂY DỰNG MODEL TRẢ LỜI CÂU HỎI
##Model_NLP
#List key gg sheet
def next_key(index,list_key):
    if index < len(list_key)-1:
        key=list_key[index]
        index=index+1
        return key,index
    else:
        index=0
        key_next=list_key[index]
        index=index+1
        return key_next,index
    
 
#processing file doc
def read_docx(file_path,list_key):
    # Tạo một đối tượng Document từ tệp DOCX
    doc = Document(file_path)
    file_path_old=url(r'folder/content.docx')
    print(file_path_old)
    #file_path_old='D:/AI_test_def_and_ggsheet/folder/content.docx'
    doc_old=Document(file_path_old)
    # Khởi tạo biến chứa nội dung
    content = ""
    content_old=''
    index_key_next=0
    index_paragraph=0
    index_name=0
    your_openai_key,index_key_next = next_key(index_key_next,list_key)
    client = openai.OpenAI(api_key=your_openai_key)
    list_content=doc.paragraphs
    list_content_old=doc_old.paragraphs
    list_content_size=len(list_content)
    for paragraph in list_content:
        content += paragraph.text + "\n"
    for paragraph in list_content_old:
        content_old += paragraph.text + "\n"
    if content != content_old:
        remove_data(r'folder\speech_content_folder')
        while index_paragraph < list_content_size:
            if list_content[index_paragraph].text !='':
                #Tách content thành các đoạn nhỏ rồi chuyển sang speech
                try:
                    text_to_speech(list_content[index_paragraph].text,index_name,client,folder = url(r'folder/speech_content_folder'))
                    index_paragraph=index_paragraph+1
                    index_name=index_name+1
                except openai.RateLimitError:
                    your_openai_key,index_key_next=next_key(index_key_next,list_key)
                    print(your_openai_key)
                    client = openai.OpenAI(api_key=your_openai_key)
                except openai.AuthenticationError:
                    print('key lỗi',your_openai_key)
                    index_key_next=index_key_next-1
                    list_key.pop(index_key_next)
                    url_api_sheet=url(r'api_sheet.json')
                    gc = gspread.service_account(filename=url_api_sheet )
                    wks = gc.open("API_chatgpt").sheet1
                    cell = wks.find(your_openai_key)
                    cell='A'+str(cell.row)
                    wks.update(cell, '')
                    your_openai_key,index_key_next=next_key(index_key_next,list_key)
                    client = openai.OpenAI(api_key=your_openai_key)
                except openai.APITimeoutError:
                    time.sleep(10)
            else:
                index_paragraph=index_paragraph+1
        doc.save(file_path_old)
    return content
#Model chatgpt
def model_NLP(user, comment, content,description,client, max_tokens=150, temperature=0.5):
        messages = [
            {"role": "system", "content": f"{description} {content}"},
            {"role": "user", "content": f"Hãy trả lời câu hỏi: {comment}"},
            #{"role": "assistant", "content": f"Xin chào {user} và câu hỏi của bạn: {comment}"}
            ]
        api_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature,)
        answer = api_response.choices[0].message.content
        if answer[-1] != '.':
            answer = ".".join(answer.split(sep=".")[:-1])
        if answer != '':
            answer= 'Xin chào '+str(user)+", "+str(answer)
        return answer

#XÂY DỰNG BỘ CHUYỂN VĂN BẢN SANG ÂM THANH
##Text-to-speech
def text_to_speech(answer,name_file,client,folder= url(r'speech_folder')):
    if answer !='':
        response = client.audio.speech.create(
        model="tts-1",
        voice="nova", # other voices: 'echo', 'fable', 'onyx', 'nova', 'shimmer','alloy'
        input=answer)
        name_file_speech=str(name_file) +'.mp3'
        output_video_path = os.path.join(folder, name_file_speech)
        response.stream_to_file(output_video_path)

## speech_to_video
#Hàm chỉ đường dẫn tiếp theo của file trong tệp speech_content_folder
def urf_file_speech_content_continue(index,index_content):
    url_speech_content_folder=url(r'folder/speech_content_folder')
    url_video_content_folder=url(r'folder/video_content_folder')
    speech_files_list = [os.path.join(url_speech_content_folder, file) for file in os.listdir(url_speech_content_folder)]
    speech_files_size=len(speech_files_list)
    if index_content < speech_files_size-1:
        index_content_new=index_content+1
        name_file= str(index_content_new)+'.mp3'
        name_file_video=str(index_content_new)+'.mp4'
        url_file_speech=os.path.join(url_speech_content_folder, name_file)
        url_file_video=os.path.join(url_video_content_folder, name_file_video)
        return url_file_speech,url_file_video,index,index_content_new
    else:
        index_content_new=0
        name_file= str(index_content_new)+'.mp3'
        url_file_speech=os.path.join(url_speech_content_folder, name_file)
        name_file_video=str(index_content_new)+'.mp4'
        url_file_video=os.path.join(url_video_content_folder, name_file_video)
        return url_file_speech,url_file_video,index,index_content_new
    
#Hàm đưa ra đường dẫn tiếp theo cần thực hiện  
def urf_file_speech_continue(index,index_content):
    url_speech_folder=url(r'speech_folder')
    print(url_speech_folder)
    speech_files_list = [os.path.join(url_speech_folder, file) for file in os.listdir(url_speech_folder)]
    speech_files_size=len(speech_files_list)
    if speech_files_size !=0:
        if index < speech_files_size-1:
            index_new=index+1
            name_file= str(index_new)+'.mp3'
            url_file=os.path.join(url_speech_folder, name_file)
            tg=get_audio_time(url_file)
            number_frame=int(tg*30)
            list_img=[]
            folder_img=url(r'folder/data_img')
            i=0
            while i < number_frame:
                if i <= 263:
                    name_img=str(i)+'.png'
                    url_img =folder_img+'/'+name_img
                    list_img.append(url_img)
                    i=i+1
                else:
                    number_frame=number_frame-i
                    i=0
                    name_img=str(i)+'.png'
                    name_img=str(i)+'.png'
                    url_img =folder_img+'/'+name_img
                    list_img.append(url_img)
                    i=i+1
            return url_file,list_img,index_new,index_content
        else:
            return urf_file_speech_content_continue(index,index_content)
    else:
        return urf_file_speech_content_continue(index,index_content)

# XÓA
##hàm xóa dữ liệu khi chương trình kết thúc  
def remove_data(name_folder): 
    url_folder = Path(__file__).parent
    url = url_folder / Path(name_folder)
    #xóa file trong tệp speech_folder
    speech_files_list = [os.path.join(url,file) for file in os.listdir(url)]
    for speech_file in speech_files_list:
        os.remove(speech_file)
# LUỒNG
##LUỒNG 1
def process_thread_1(id):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    loop.create_task(crawling(id))
    loop.run_forever()
    
##LUỒNG 2
#content là biến golbal được khởi tạo là rỗng nhưng sẽ được dán nội dung vào khi thực thi hàm import_file trong button.py
def process_thread_2(content,description,list_key):
    global list_data
    index_name=0
    location=0
    index_key_next=0
    your_openai_key,index_key_next = next_key(index_key_next,list_key)
    client = openai.OpenAI(api_key=your_openai_key)
    while True:
            if len(list_data)!=0 and location != len(list_data):
                #dataset.to_csv('data.csv',index=False)
                while location < len(list_data) :
                    try:
                        answer =model_NLP(list_data[location][0],list_data[location][1],content,description,client)
                        if answer !='':
                            text_to_speech(answer,index_name,client)
                            print(list_data[location][0],", ",list_data[location][1])
                            print(answer)
                            index_name = index_name +1
                            location=location+1
                    except openai.RateLimitError:
                        your_openai_key,index_key_next=next_key(index_key_next,list_key)
                        print(your_openai_key)
                        client = openai.OpenAI(api_key=your_openai_key)
                    except openai.AuthenticationError:
                        print('key lỗi',your_openai_key)
                        index_key_next=index_key_next-1
                        list_key.pop(index_key_next)
                        gc = gspread.service_account(filename=url(r'api_sheet.json'))
                        wks = gc.open("API_chatgpt").sheet1
                        cell = wks.find(your_openai_key)
                        cell='A'+str(cell.row)
                        wks.update(cell, '')
                        your_openai_key,index_key_next=next_key(index_key_next,list_key)
                        client = openai.OpenAI(api_key=your_openai_key)
                    except openai.APITimeoutError:
                        time.sleep(10)
                location=len(list_data)
            '''elif status==False:
                location=len(list_data)
                print('Đang dừng trả lời câu hỏi')'''
##LUỒNG 3

